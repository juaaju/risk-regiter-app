import os
import uuid
import shutil
import logging
from datetime import datetime
from typing import List, Optional
from enum import Enum

import lmstudio as lms
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import psycopg2
from psycopg2.extras import Json, DictCursor
import asyncpg
import PyPDF2
import docx
import pandas as pd
from sqlalchemy import create_engine
import numpy as np
import traceback

# Konfigurasi logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("api.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("rag-api")

# Inisialisasi model embedding
try:
    logger.info("Mencoba inisialisasi model embedding...")
    model = lms.embedding_model("nomic-embed-text-v1.5")
    logger.info("Model embedding berhasil diinisialisasi.")
except Exception as e:
    logger.error(f"Error saat inisialisasi model embedding: {str(e)}")
    logger.error(traceback.format_exc())
    # Jangan crash aplikasi saat ini, berikan pesan yang jelas

app = FastAPI(title="Document Embedding API")

# Konfigurasi CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Izinkan semua origin - untuk development
    allow_credentials=True,
    allow_methods=["*"],  # Izinkan semua HTTP methods
    allow_headers=["*"],  # Izinkan semua headers
)

# Konfigurasi database
DB_CONFIG = {
    "host": "localhost",
    "port": 5432,
    "database": "risk_management",
    "user": "postgres",
    "password": "toor"
}

# Lokasi penyimpanan file
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Enum untuk kategori dokumen
class DocumentCategory(str, Enum):
    standard_safety = "standard_safety"
    general_document = "general_document"

# Model untuk respons
class DocumentResponse(BaseModel):
    id: int
    filename: str
    original_filename: str
    category: str
    file_size: int
    upload_date: datetime
    processing_status: str

# Fungsi untuk koneksi ke database
async def get_db_connection():
    try:
        logger.debug("Mencoba membuat koneksi asyncpg...")
        conn = await asyncpg.connect(
            user=DB_CONFIG["user"],
            password=DB_CONFIG["password"],
            database=DB_CONFIG["database"],
            host=DB_CONFIG["host"],
            port=DB_CONFIG["port"]
        )
        logger.debug("Koneksi asyncpg berhasil dibuat.")
        return conn
    except Exception as e:
        logger.error(f"Error membuat koneksi asyncpg: {str(e)}")
        logger.error(traceback.format_exc())
        raise

def get_psycopg2_connection():
    try:
        logger.debug("Mencoba membuat koneksi psycopg2...")
        conn = psycopg2.connect(
            user=DB_CONFIG["user"],
            password=DB_CONFIG["password"],
            dbname=DB_CONFIG["database"],
            host=DB_CONFIG["host"],
            port=DB_CONFIG["port"]
        )
        logger.debug("Koneksi psycopg2 berhasil dibuat.")
        return conn
    except Exception as e:
        logger.error(f"Error membuat koneksi psycopg2: {str(e)}")
        logger.error(traceback.format_exc())
        raise

# Fungsi untuk ekstraksi teks dari berbagai format file
def extract_text(file_path):
    file_extension = file_path.split('.')[-1].lower()
    logger.info(f"Mengekstrak teks dari file: {file_path} dengan ekstensi: {file_extension}")
    
    try:
        if file_extension == 'pdf':
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                logger.info(f"PDF memiliki {len(pdf_reader.pages)} halaman")
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    logger.debug(f"Halaman {page_num+1}: Diproses {len(page_text)} karakter")
            logger.info(f"Ekstraksi PDF selesai. Total karakter: {len(text)}")
            return text
        
        elif file_extension in ['docx', 'doc']:
            try:
                logger.info("Mencoba mengekstrak dokumen Word...")
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                logger.info(f"Ekstraksi Word selesai. Total paragraf: {len(doc.paragraphs)}, Karakter: {len(text)}")
                return text
            except Exception as e:
                logger.error(f"Error saat mengekstrak dokumen Word: {str(e)}")
                logger.error(traceback.format_exc())
                return ""
        
        elif file_extension == 'txt':
            logger.info("Mengekstrak file teks...")
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
                logger.info(f"Ekstraksi teks selesai. Total karakter: {len(text)}")
                return text
        
        elif file_extension in ['csv', 'xlsx', 'xls']:
            try:
                logger.info(f"Mengekstrak file {file_extension}...")
                if file_extension == 'csv':
                    df = pd.read_csv(file_path)
                else:
                    df = pd.read_excel(file_path)
                logger.info(f"File berhasil dibaca. Jumlah baris: {len(df)}, Kolom: {len(df.columns)}")
                text = df.to_string()
                logger.info(f"Konversi dataframe ke teks: {len(text)} karakter")
                return text
            except Exception as e:
                logger.error(f"Error saat mengekstrak file {file_extension}: {str(e)}")
                logger.error(traceback.format_exc())
                return ""
        
        else:
            logger.warning(f"Ekstensi file tidak didukung: {file_extension}")
            return ""
            
    except Exception as e:
        logger.error(f"Error tidak terduga saat ekstraksi teks: {str(e)}")
        logger.error(traceback.format_exc())
        return ""

# Fungsi untuk membagi teks menjadi chunks
def chunk_text(text, chunk_size=1000, overlap=200):
    logger.info(f"Membagi teks menjadi chunks. Panjang teks: {len(text)}, Ukuran chunk: {chunk_size}, Overlap: {overlap}")
    
    if not text:
        logger.warning("Teks kosong, tidak ada chunk yang dibuat")
        return []
    
    # Bagi berdasarkan paragraf dulu
    paragraphs = text.split('\n')
    logger.debug(f"Teks dibagi menjadi {len(paragraphs)} paragraf")
    
    chunks = []
    current_chunk = ""
    
    for i, paragraph in enumerate(paragraphs):
        if len(current_chunk) + len(paragraph) <= chunk_size:
            current_chunk += paragraph + "\n"
        else:
            # Simpan chunk saat ini
            if current_chunk:
                chunks.append(current_chunk.strip())
                logger.debug(f"Chunk {len(chunks)} dibuat dengan {len(current_chunk)} karakter")
            
            # Mulai chunk baru, mungkin dengan overlap
            words = current_chunk.split()
            if len(words) > overlap:
                overlap_text = " ".join(words[-overlap:])
                current_chunk = overlap_text + "\n" + paragraph + "\n"
                logger.debug(f"Overlap: {len(overlap_text)} karakter ditambahkan ke chunk baru")
            else:
                current_chunk = paragraph + "\n"
    
    # Tambahkan chunk terakhir jika ada
    if current_chunk:
        chunks.append(current_chunk.strip())
        logger.debug(f"Chunk terakhir ({len(chunks)}) dibuat dengan {len(current_chunk)} karakter")
    
    logger.info(f"Total {len(chunks)} chunks dibuat")
    return chunks

# Fungsi untuk pemrosesan dokumen dan pembuatan embedding
# Fungsi untuk pemrosesan dokumen dan pembuatan embedding
async def process_document(document_id, file_path):
    logger.info(f"Memulai pemrosesan dokumen ID {document_id}: {file_path}")
    conn = None
    cursor = None
    
    try:
        conn = get_psycopg2_connection()
        cursor = conn.cursor()
        
        # Update status pemrosesan
        cursor.execute(
            "UPDATE documents SET processing_status = 'processing' WHERE id = %s",
            (document_id,)
        )
        conn.commit()
        logger.info(f"Status dokumen {document_id} diubah menjadi 'processing'")
        
        # Ekstrak teks dari dokumen
        logger.info(f"Memulai ekstraksi teks dari dokumen {document_id}")
        extracted_text = extract_text(file_path)  # Gunakan nama variabel yang berbeda
        
        if not extracted_text:
            # Update status jika ekstraksi gagal
            logger.error(f"Ekstraksi teks gagal untuk dokumen {document_id}")
            cursor.execute(
                "UPDATE documents SET processing_status = 'failed', error_message = 'Tidak dapat mengekstrak teks dari dokumen' WHERE id = %s",
                (document_id,)
            )
            conn.commit()
            return
        
        logger.info(f"Ekstraksi teks berhasil untuk dokumen {document_id}. Panjang teks: {len(extracted_text)}")
        
        # Chunk teks menjadi bagian-bagian kecil
        logger.info(f"Memulai chunking teks untuk dokumen {document_id}")
        text_chunks = chunk_text(extracted_text)  # Gunakan nama variabel yang berbeda
        
        if not text_chunks:
            logger.error(f"Chunking gagal untuk dokumen {document_id}")
            cursor.execute(
                "UPDATE documents SET processing_status = 'failed', error_message = 'Gagal membagi teks menjadi chunks' WHERE id = %s",
                (document_id,)
            )
            conn.commit()
            return
        
        logger.info(f"Memulai pembuatan embedding untuk {len(text_chunks)} chunks dari dokumen {document_id}")
        
        # Buat embedding untuk setiap chunk
        for i, chunk_content in enumerate(text_chunks):  # Gunakan nama variabel yang berbeda
            try:
                logger.debug(f"Membuat embedding untuk chunk {i+1}/{len(text_chunks)} dari dokumen {document_id}")
                # Buat embedding menggunakan model
                embedding = model.embed(chunk_content)
                
                # Metadata untuk chunk
                metadata = {
                    'chunk_index': i,
                    'char_length': len(chunk_content)
                }
                
                logger.debug(f"Menyimpan chunk {i+1} dan embedding ke database")
                # Simpan chunk dan embedding ke database
                
                with conn.cursor() as chunk_cursor:
                    chunk_cursor.execute(
                        """
                        INSERT INTO document_chunks (document_id, chunk_index, content, metadata, embedding)
                        VALUES (%s, %s, %s, %s, %s)
                        """,
                        (document_id, i, chunk_content, Json(metadata), embedding)
                    )
                    conn.commit()  # Commit setiap insert
                    
            except Exception as e:
                conn.rollback()  # Rollback jika terjadi kesalahan
                logger.error(f"Error saat memproses chunk {i+1}: {str(e)}")
                # Lanjutkan ke chunk berikutnya
        
        # Update status pemrosesan menjadi completed
        logger.info(f"Pemrosesan dokumen {document_id} selesai, mengupdate status")
        cursor.execute(
            "UPDATE documents SET processing_status = 'completed', last_processed_date = CURRENT_TIMESTAMP WHERE id = %s",
            (document_id,)
        )
        conn.commit()
        logger.info(f"Status dokumen {document_id} diubah menjadi 'completed'")
        
    except Exception as e:
        logger.error(f"Error tidak terduga saat pemrosesan dokumen {document_id}: {str(e)}")
        logger.error(traceback.format_exc())
        # Update status jika terjadi error
        if cursor and conn:
            try:
                cursor.execute(
                    "UPDATE documents SET processing_status = 'failed', error_message = %s WHERE id = %s",
                    (str(e), document_id)
                )
                conn.commit()
                logger.info(f"Status dokumen {document_id} diubah menjadi 'failed' dengan pesan error")
            except Exception as update_err:
                logger.error(f"Gagal mengupdate status error untuk dokumen {document_id}: {str(update_err)}")
        
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()
        logger.info(f"Koneksi database ditutup untuk dokumen {document_id}")

# Global error handler
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logger.error(f"Global error handler caught: {str(exc)}")
    logger.error(traceback.format_exc())
    return JSONResponse(
        status_code=500,
        content={"detail": f"Internal Server Error: {str(exc)}"}
    )

# Endpoint untuk mengunggah file
@app.post("/api/documents/upload", response_model=DocumentResponse)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    category: DocumentCategory = Form(...),
    description: Optional[str] = Form(None)
):
    logger.info(f"Menerima request upload file: {file.filename}, kategori: {category}")
    
    # Buat nama file unik
    file_extension = file.filename.split('.')[-1]
    unique_filename = f"{uuid.uuid4()}.{file_extension}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    logger.info(f"Menyimpan file ke: {file_path}")
    # Simpan file ke disk
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        logger.info(f"File berhasil disimpan di: {file_path}")
    except Exception as e:
        logger.error(f"Error saat menyimpan file: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Gagal menyimpan file: {str(e)}")
    
    # Dapatkan ukuran file
    file_size = os.path.getsize(file_path)
    logger.info(f"Ukuran file: {file_size} bytes")
    
    # Simpan metadata ke database
    conn = None
    try:
        logger.info("Menyimpan metadata dokumen ke database")
        conn = await get_db_connection()
        document_id = await conn.fetchval(
            """
            INSERT INTO documents 
            (filename, original_filename, file_path, file_type, file_size, category, description, created_by)
            VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
            RETURNING id
            """,
            unique_filename, file.filename, file_path, file.content_type, file_size, category, description, "system"
        )
        logger.info(f"Dokumen berhasil disimpan dengan ID: {document_id}")
        
        # Dapatkan data dokumen untuk respons
        document = await conn.fetchrow(
            "SELECT id, filename, original_filename, category, file_size, upload_date, processing_status FROM documents WHERE id = $1",
            document_id
        )
        
        # Mulai proses embedding di background
        logger.info(f"Memulai proses embedding di background untuk dokumen ID: {document_id}")
        background_tasks.add_task(process_document, document_id, file_path)
        
        result = {
            "id": document["id"],
            "filename": document["filename"],
            "original_filename": document["original_filename"],
            "category": document["category"],
            "file_size": document["file_size"],
            "upload_date": document["upload_date"],
            "processing_status": document["processing_status"]
        }
        logger.info(f"Mengembalikan response untuk dokumen ID: {document_id}")
        return result
    
    except Exception as e:
        logger.error(f"Error saat menyimpan metadata dokumen: {str(e)}")
        logger.error(traceback.format_exc())
        # Hapus file jika gagal menyimpan ke database
        try:
            os.remove(file_path)
            logger.info(f"File dihapus karena error: {file_path}")
        except:
            logger.error(f"Gagal menghapus file setelah error: {file_path}")
        
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        if conn:
            await conn.close()
            logger.debug("Koneksi database ditutup")

# Endpoint untuk mendapatkan semua dokumen
@app.get("/api/documents", response_model=List[DocumentResponse])
async def get_documents(category: Optional[DocumentCategory] = None):
    logger.info(f"Request untuk mendapatkan semua dokumen. Filter kategori: {category}")
    conn = await get_db_connection()
    try:
        if category:
            logger.info(f"Mengambil dokumen dengan kategori: {category}")
            rows = await conn.fetch(
                """
                SELECT id, filename, original_filename, category, file_size, upload_date, processing_status 
                FROM documents WHERE category = $1 ORDER BY upload_date DESC
                """,
                category
            )
        else:
            logger.info("Mengambil semua dokumen tanpa filter")
            rows = await conn.fetch(
                """
                SELECT id, filename, original_filename, category, file_size, upload_date, processing_status 
                FROM documents ORDER BY upload_date DESC
                """
            )
        
        logger.info(f"Mengembalikan {len(rows)} dokumen")
        return [dict(row) for row in rows]
    
    except Exception as e:
        logger.error(f"Error saat mengambil dokumen: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        await conn.close()
        logger.debug("Koneksi database ditutup")

# Endpoint untuk mendapatkan dokumen berdasarkan ID
@app.get("/api/documents/{document_id}", response_model=DocumentResponse)
async def get_document(document_id: int):
    logger.info(f"Request untuk mendapatkan dokumen dengan ID: {document_id}")
    conn = await get_db_connection()
    try:
        document = await conn.fetchrow(
            """
            SELECT id, filename, original_filename, category, file_size, upload_date, processing_status 
            FROM documents WHERE id = $1
            """,
            document_id
        )
        
        if not document:
            logger.warning(f"Dokumen dengan ID {document_id} tidak ditemukan")
            raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan")
        
        logger.info(f"Dokumen dengan ID {document_id} berhasil ditemukan")
        return dict(document)
    
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        logger.error(f"Error saat mengambil dokumen ID {document_id}: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        await conn.close()
        logger.debug("Koneksi database ditutup")

# Endpoint untuk menghapus dokumen
@app.delete("/api/documents/{document_id}")
async def delete_document(document_id: int):
    logger.info(f"Request untuk menghapus dokumen dengan ID: {document_id}")
    conn = await get_db_connection()
    try:
        # Dapatkan informasi file sebelum dihapus
        file_info = await conn.fetchrow(
            "SELECT file_path FROM documents WHERE id = $1",
            document_id
        )
        
        if not file_info:
            logger.warning(f"Dokumen dengan ID {document_id} tidak ditemukan untuk dihapus")
            raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan")
        
        logger.info(f"Menghapus chunks terkait untuk dokumen ID {document_id}")
        # Hapus chunks terkait dari database
        await conn.execute(
            "DELETE FROM document_chunks WHERE document_id = $1",
            document_id
        )
        
        logger.info(f"Menghapus dokumen ID {document_id} dari database")
        # Hapus dokumen dari database
        await conn.execute(
            "DELETE FROM documents WHERE id = $1",
            document_id
        )
        
        # Hapus file fisik
        file_path = file_info["file_path"]
        if os.path.exists(file_path):
            logger.info(f"Menghapus file fisik: {file_path}")
            os.remove(file_path)
            logger.info(f"File fisik berhasil dihapus: {file_path}")
        else:
            logger.warning(f"File fisik tidak ditemukan untuk dihapus: {file_path}")
        
        logger.info(f"Dokumen ID {document_id} berhasil dihapus")
        return {"message": "Dokumen berhasil dihapus"}
    
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        logger.error(f"Error saat menghapus dokumen ID {document_id}: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        await conn.close()
        logger.debug("Koneksi database ditutup")

# Endpoint untuk search berdasarkan vektor similarity
@app.get("/api/search")
async def search_documents(query: str, category: Optional[DocumentCategory] = None, max_results: int = 5):
    logger.info(f"Request pencarian. Query: '{query}', Kategori: {category}, Max results: {max_results}")
    
    if not query.strip():
        logger.warning("Query pencarian kosong")
        return {"results": []}
    
    # Buat embedding dari query
    try:
        logger.info("Membuat embedding dari query pencarian")
        query_embedding = model.embed(query)
        logger.debug(f"Embedding query berhasil dibuat dengan dimensi {len(query_embedding)}")
    except Exception as e:
        logger.error(f"Error saat membuat embedding query: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Gagal membuat embedding query: {str(e)}")
    
    conn = None
    cursor = None
    try:
        conn = get_psycopg2_connection()
        cursor = conn.cursor(cursor_factory=DictCursor)
        
        # Simpan riwayat kueri
        logger.info("Menyimpan riwayat kueri ke database")
        cursor.execute(
            """
            INSERT INTO query_history (query_text, query_embedding, num_results)
            VALUES (%s, %s, %s) RETURNING id
            """,
            (query, query_embedding, max_results)
        )
        query_id = cursor.fetchone()[0]
        conn.commit()
        logger.info(f"Riwayat kueri disimpan dengan ID: {query_id}")
        
        # Cari dokumen berdasarkan similarity
        logger.info("Mencari dokumen berdasarkan similarity vektor")
        if category:
            logger.info(f"Pencarian dengan filter kategori: {category}")
            cursor.execute(
                """
                SELECT * FROM search_similar_chunks(%s, 0.5, %s, %s)
                """,
                (query_embedding, max_results, category)
            )
        else:
            logger.info("Pencarian tanpa filter kategori")
            cursor.execute(
                """
                SELECT * FROM search_similar_chunks(%s, 0.5, %s)
                """,
                (query_embedding, max_results)
            )
        
        results = cursor.fetchall()
        logger.info(f"Pencarian mengembalikan {len(results)} hasil")
        
        # Format hasil
        formatted_results = []
        for row in results:
            logger.debug(f"Hasil: chunk_id={row['chunk_id']}, dokumen={row['filename']}, similarity={row['similarity']}")
            formatted_results.append({
                "chunk_id": row["chunk_id"],
                "document_id": row["document_id"],
                "filename": row["filename"],
                "category": row["category"],
                "content": row["content"],
                "similarity": float(row["similarity"])
            })
            
            # Catat akses dokumen
            logger.debug(f"Mencatat akses dokumen ID {row['document_id']} dari kueri ID {query_id}")
            cursor.execute(
                """
                INSERT INTO document_access_logs (document_id, access_type, query_id)
                VALUES (%s, %s, %s)
                """,
                (row["document_id"], "search", query_id)
            )
        
        conn.commit()
        logger.info(f"Mengembalikan {len(formatted_results)} hasil yang diformat")
        return {"results": formatted_results}
    
    except Exception as e:
        logger.error(f"Error saat pencarian dokumen: {str(e)}")
        logger.error(traceback.format_exc())
        if conn:
            conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()
        logger.debug("Koneksi database ditutup")

# Jalankan server dengan uvicorn
if __name__ == "__main__":
    import uvicorn
    logger.info("Memulai server API...")
    uvicorn.run(app, host="0.0.0.0", port=8005)